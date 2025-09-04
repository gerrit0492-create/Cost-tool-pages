# pages/18_Offerte_DOCX.py  (complete, met fix voor st.dataframe)
import os, io, json, math
from datetime import date
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.set_page_config(page_title="Offerte export (DOCX)", page_icon="🧾", layout="wide")
st.title("🧾 Offerte export (DOCX) - met logo, btw en nette opmaak")

# ---- Kopgegevens
col1, col2 = st.columns(2)
client_name    = col1.text_input("Klantnaam", "Acme BV")
client_contact = col1.text_input("Contact", "J. Janssen")
client_email   = col1.text_input("Email", "sales@acme.nl")
project_code   = col2.text_input("Projectcode", "RFQ-2025-001")
lead_weeks     = col2.number_input("Levertijd (weken)", 1, 52, 4)
vat_pct        = col2.number_input("BTW %", 0, 30, 21)

# ---- Logo (optioneel)
logo_file = st.file_uploader("Logo (PNG/JPG, optioneel; komt in de kop)", type=["png","jpg","jpeg"])

# ---- Check data-bestanden
missing = []
for p in ["data/bom_current.json", "data/material_prices.csv", "data/labor_rates.csv"]:
    if not os.path.exists(p):
        missing.append(p)
if missing:
    st.error("Ontbrekend voor DOCX-offerte:\n- " + "\n".join(missing))
    st.stop()

# ---- Data laden
bom = json.load(open("data/bom_current.json"))
items = bom.get("bom", [])
assembly = bom.get("assembly", {"name":"", "qty":1})
df_prices = pd.read_csv("data/material_prices.csv")
df_rates  = pd.read_csv("data/labor_rates.csv")

# ---- Helpers
DENSITY_KG_PER_MM3 = {"stainless":7.9e-6,"duplex":7.8e-6,"aluminum":2.7e-6,"carbon_steel":7.85e-6}

def infer_family_from_grade(grade:str)->str:
    g=(grade or "").lower()
    if "1.44" in g or "2205" in g or "s31803" in g or "s32205" in g: return "duplex"
    if any(k in g for k in ["304","316","1.43","1.45"]): return "stainless"
    if any(k in g for k in ["6082","6061","5754","1050","alu","aluminium"]): return "aluminum"
    return "carbon_steel"

def mass_kg(part):
    fam=(part.get("material_family") or infer_family_from_grade(part.get("material_grade",""))).lower()
    rho=DENSITY_KG_PER_MM3.get(fam,7.85e-6)
    form=(part.get("form") or "").lower()
    t=part.get("thickness_mm") or 0; L=part.get("length_mm") or 0
    W=part.get("width_mm") or 0; d=part.get("diameter_mm") or 0
    if form in ("sheet","plate") and t and L and W:
        vol=float(t)*float(L)*float(W)
    elif form in ("bar","staf","round","rod") and d and L:
        vol=math.pi*(float(d)/2)**2*float(L)
    else:
        vol=float(t or 0)*float(L or 0)*float(W or 0)  # fallback
    return max(0.0, vol)*rho

def latest_material_price(df, grade, region="EU", unit="€/kg"):
    if df.empty: return 0.0
    d=df.copy()
    d=d[d["grade"].astype(str).str.lower()==str(grade).lower()]
    if region: d=d[d["region"].astype(str).str.upper()==str(region).upper()]
    if unit:   d=d[d["unit"].astype(str)==unit]
    if d.empty: return 0.0
    d["as_of_date"]=pd.to_datetime(d["as_of_date"], errors="coerce")
    d=d.sort_values("as_of_date")
    return float(d.iloc[-1]["price"])

def midpoint_rate(df, process_kw, country="Netherlands"):
    """Zoek midden van min/max voor proces; valt terug per keyword match."""
    if df.empty: return 1.0/60.0
    d=df.copy()
    d=d[d["process"].astype(str).str.lower().str.contains(process_kw.lower())]
    if not d.empty:
        d2=d[d["country"].astype(str).str.lower()==country.lower()]
        if not d2.empty: d=d2
    if d.empty: return 1.0/60.0
    d["as_of_date"]=pd.to_datetime(d["as_of_date"], errors="coerce")
    r=d.sort_values("as_of_date").iloc[-1]
    return float((r["rate_min"]+r["rate_max"])/2.0)/60.0  # €/min

def map_rate_for_process(df_rates, proc_name:str):
    """Slimmere mapping: direct proces als het bestaat, anders beste alternatief."""
    p=proc_name.strip().lower()
    direct_map = {
        "laser": ["laser","lasersnijden"],
        "bend": ["bend","buigen","kantbank","press brake"],
        "tig": ["tig","tig welding","lassen tig"],
        "cnc_mill": ["cnc mill","cnc milling","frezen"],
        "cnc_turn": ["cnc turn","cnc turning","draaien"]
    }
    for key, kws in direct_map.items():
        if p==key or any(p==k or p in k for k in kws):
            for kw in [p] + kws:
                rate = midpoint_rate(df_rates, kw)
                if rate>0.02:
                    return rate
    if "laser" in p:   return midpoint_rate(df_rates, "cnc milling")
    if "bend" in p:    return midpoint_rate(df_rates, "cnc milling")
    if "tig" in p:     return midpoint_rate(df_rates, "tig")
    if "mill" in p:    return midpoint_rate(df_rates, "cnc milling")
    if "turn" in p:    return midpoint_rate(df_rates, "cnc turning")
    return midpoint_rate(df_rates, p)

def est_minutes(part, proc):
    q=max(1, int(part.get("qty",1)))
    base={"laser":(5,0.43),"bend":(8,0.50),"tig":(10,0.60),"cnc_mill":(12,1.20),"cnc_turn":(10,1.00)}
    key = proc.strip().lower()
    setup,cycle = base.get(key, (5,0.30))
    return (setup/q)+cycle

# ---- Reken per item
rows=[]
for p in items:
    grade=p.get("material_grade","")
    fam  =p.get("material_family") or infer_family_from_grade(grade)
    m_kg =mass_kg(p)
    eur_per_kg = latest_material_price(df_prices, grade)
    mat_eur_pc = m_kg * eur_per_kg

    procs=[x.strip() for x in (p.get("processes") or [])]
    proc_detail=[]; proc_cost_pc=0.0
    for proc in procs:
        norm = proc.upper() if proc.upper()=="TIG" else proc.lower()
        rate = map_rate_for_process(df_rates, norm)
        minutes=est_minutes(p, "TIG" if norm=="TIG" else norm)
        cost_pc=minutes*rate
        proc_cost_pc+=cost_pc
        proc_detail.append((proc, round(rate,2), round(minutes,2), round(cost_pc,2)))

    total_pc=mat_eur_pc+proc_cost_pc
    rows.append({
        "item": p.get("item_code","?"),
        "qty": int(p.get("qty",1)),
        "grade": grade,
        "family": fam,
        "mass": round(m_kg,4),
        "eur_kg": round(eur_per_kg,4),
        "mat_pc": round(mat_eur_pc,2),
        "proc_pc": round(proc_cost_pc,2),
        "tot_pc": round(total_pc,2),
        "proc_detail": proc_detail
    })

df=pd.DataFrame(rows)

# ---- Samenvatting & tabel in de app (zonder complexe proc_detail kolom)
total_excl = float((df["tot_pc"] * df["qty"]).sum()) if not df.empty else 0.0
total_incl = total_excl * (1 + vat_pct/100.0)

st.subheader("Samenvatting")
c1,c2,c3=st.columns(3)
c1.metric("Totaal excl. btw", f"€ {total_excl:,.2f}")
c2.metric("BTW", f"{vat_pct}%")
c3.metric("Totaal incl. btw", f"€ {total_incl:,.2f}")

cols = ["item","qty","grade","family","mass","eur_kg","mat_pc","proc_pc","tot_pc"]
df_display = df[cols].copy()
for c in ["qty","mass","eur_kg","mat_pc","proc_pc","tot_pc"]:
    df_display[c] = pd.to_numeric(df_display[c], errors="coerce")

st.dataframe(df_display, use_container_width=True)

# Download schone CSV (zonder proc_detail)
st.download_button(
    "⬇️ Download resultaten.csv",
    df_display.to_csv(index=False),
    "offerte_resultaten.csv",
    "text/csv"
)

# ---- DOCX opbouwen
doc=Document()

# Kop met logo en titel
sec = doc.sections[0]
header = sec.header.paragraphs[0]
if logo_file:
    run = header.add_run()
    run.add_picture(logo_file, width=Inches(1.2))
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.add_run("  ")
else:
    header.add_run("")

title = doc.add_heading(f"Offer {project_code} - {client_name}", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

meta = doc.add_paragraph()
meta_run = meta.add_run(f"Date: {date.today().isoformat()}    ")
meta_run.bold = True
meta.add_run(f"Contact: {client_contact} | {client_email}")

doc.add_heading("Scope", level=2)
doc.add_paragraph(f"Assembly: {assembly.get('name','')} - quantity {assembly.get('qty',1)}.")

doc.add_heading("Cost breakdown (per piece, EUR)", level=2)
table=doc.add_table(rows=1, cols=7)
table.style = "Light List Accent 1"
hdr=table.rows[0].cells
hdr[0].text="Item"; hdr[1].text="Grade"; hdr[2].text="Mass (kg/pc)"
hdr[3].text="€/kg"; hdr[4].text="Material €/pc"; hdr[5].text="Proc €/pc"; hdr[6].text="Total €/pc"

for r in rows:
    row=table.add_row().cells
    row[0].text=str(r["item"])
    row[1].text=str(r["grade"])
    row[2].text=f'{r["mass"]:.4f}'
    row[3].text=f'{r["eur_kg"]:.4f}'
    row[4].text=f'{r["mat_pc"]:.2f}'
    row[5].text=f'{r["proc_pc"]:.2f}'
    row[6].text=f'{r["tot_pc"]:.2f}'

doc.add_paragraph().add_run(f"Assembly total (qty × Total €/pc): € {total_excl:,.2f} excl. btw").bold=True
doc.add_paragraph(f"VAT {vat_pct}% → Total incl. VAT: € {total_incl:,.2f}")

doc.add_heading("Process detail", level=2)
for r in rows:
    doc.add_heading(r["item"], level=3)
    if not r["proc_detail"]:
        doc.add_paragraph("No process cost lines.")
        continue
    t=doc.add_table(rows=1, cols=4); t.style = "Light Grid"
    th=t.rows[0].cells
    th[0].text="Process"; th[1].text="Rate (€/min)"; th[2].text="Minutes"; th[3].text="Cost €/pc"
    for (proc, rate, mins, cost) in r["proc_detail"]:
        rw=t.add_row().cells
        rw[0].text=str(proc); rw[1].text=f"{rate:.2f}"; rw[2].text=f"{mins:.2f}"; rw[3].text=f"{cost:.2f}"

doc.add_heading("Assumptions", level=2)
doc.add_paragraph(f"Lead time: {int(lead_weeks)} weeks after order.")
doc.add_paragraph("Incoterms: EXW.")
doc.add_paragraph("Weld quality: EN ISO 5817 level C.")
doc.add_paragraph("Scrap: 3%.")
doc.add_paragraph("Prices excl. VAT.")

# Footer (eenvoudige tekst)
footer = doc.sections[0].footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("- Page 1 -")

# ---- Download knoppen (DOCX)
buf=io.BytesIO(); doc.save(buf); buf.seek(0)
st.download_button(
    "⬇️ Download offerte.docx",
    data=buf,
    file_name=f"offerte_{project_code}.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

st.caption("Tip: voeg in data/labor_rates.csv een rij met process='laser' toe voor exact laser-tarief; anders valt hij terug op CNC milling.")
