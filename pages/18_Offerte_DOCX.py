import os, io, json
from datetime import date
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches

st.set_page_config(page_title="Offerte export (DOCX)", page_icon="🧾", layout="wide")
st.title("🧾 Offerte export (DOCX, template-loos)")

# ---- Invoer (kopgegevens)
col1, col2 = st.columns(2)
client_name    = col1.text_input("Klantnaam", "Acme BV")
client_contact = col1.text_input("Contact", "J. Janssen")
client_email   = col1.text_input("Email", "sales@acme.nl")
project_code   = col2.text_input("Projectcode", "RFQ-2025-001")
lead_weeks     = col2.number_input("Levertijd (weken)", 1, 52, 4)

# ---- Check data-bestanden
missing = []
for p in ["data/bom_current.json", "data/material_prices.csv", "data/labor_rates.csv"]:
    if not os.path.exists(p): missing.append(p)
if missing:
    st.error("Ontbrekend voor DOCX-offerte:\n- " + "\n- ".join(missing))
    st.stop()

# ---- Data laden
bom = json.load(open("data/bom_current.json"))
items = bom.get("bom", [])
assembly = bom.get("assembly", {"name":"", "qty":1})

df_prices = pd.read_csv("data/material_prices.csv")
df_rates  = pd.read_csv("data/labor_rates.csv")

# ---- Helpers (zelfde logica als Routing v1)
DENSITY_KG_PER_MM3 = {"stainless":7.9e-6,"duplex":7.8e-6,"aluminum":2.7e-6,"carbon_steel":7.85e-6}
def infer_family_from_grade(grade:str)->str:
    g=(grade or "").lower()
    if "1.44" in g or "2205" in g or "s31803" in g or "s32205" in g: return "duplex"
    if any(k in g for k in ["304","316","1.43","1.45"]): return "stainless"
    if any(k in g for k in ["6082","6061","5754","1050","alu","aluminium"]): return "aluminum"
    return "carbon_steel"

def mass_kg(part):
    fam=(part.get("material_family") or infer_family_from_grade(part.get("material_grade",""))).lower()
    rho=DENSITY_KG_PER_MM3.get(fam,7.85e-6)
    form=(part.get("form") or "").lower()
    t=part.get("thickness_mm") or 0; L=part.get("length_mm") or 0
    W=part.get("width_mm") or 0; d=part.get("diameter_mm") or 0
    if form in ("sheet","plate"):
        vol=float(t)*float(L)*float(W)
    elif form in ("bar","staf","round","rod") and d and L:
        import math; vol=math.pi*(float(d)/2)**2*float(L)
    else:
        vol=float(t)*float(L)*float(W)
    return max(0.0, vol)*rho

def latest_material_price(df, grade, region="EU", unit="€/kg"):
    d=df.copy()
    d=d[d["grade"].astype(str).str.lower()==str(grade).lower()]
    if region: d=d[d["region"].astype(str).str.upper()==str(region).upper()]
    if unit:   d=d[d["unit"].astype(str)==unit]
    if d.empty: return 0.0
    d["as_of_date"]=pd.to_datetime(d["as_of_date"], errors="coerce")
    d=d.sort_values("as_of_date")
    return float(d.iloc[-1]["price"])

def midpoint_rate(df, process, country="Netherlands"):
    d=df.copy()
    d=d[d["process"].astype(str).str.lower().str.contains(process.lower())]
    if not d.empty:
        d2=d[d["country"].astype(str).str.lower()==country.lower()]
        if not d2.empty: d=d2
    if d.empty: return 1.00/60.0
    d=d.sort_values("as_of_date")
    r=d.iloc[-1]
    return float((r["rate_min"]+r["rate_max"])/2.0)/60.0

def est_minutes(part, proc):
    q=max(1, int(part.get("qty",1)))
    base={"laser":(5,0.43),"bend":(8,0.50),"TIG":(10,0.60),"CNC_mill":(12,1.20),"CNC_turn":(10,1.00)}
    setup,cycle=base.get(proc,(5,0.30))
    return (setup/q)+cycle

# ---- Reken per item
rows=[]
for p in items:
    grade=p.get("material_grade","")
    fam  =p.get("material_family") or infer_family_from_grade(grade)
    m_kg =mass_kg(p)
    eur_per_kg = latest_material_price(df_prices, grade)
    mat_eur_pc = m_kg * eur_per_kg
    procs=[x.strip() for x in (p.get("processes") or [])]
    proc_detail=[]; proc_cost_pc=0.0
    for proc in procs:
        if proc.lower()=="laser": rate=midpoint_rate(df_rates,"CNC milling")
        elif proc.lower()=="bend": rate=midpoint_rate(df_rates,"CNC milling")
        elif proc.upper()=="TIG": rate=midpoint_rate(df_rates,"TIG")
        elif proc.lower()=="cnc_mill": rate=midpoint_rate(df_rates,"CNC milling")
        elif proc.lower()=="cnc_turn": rate=midpoint_rate(df_rates,"CNC")
        else: rate=midpoint_rate(df_rates,proc)
        minutes=est_minutes(p, "TIG" if proc=="TIG" else proc)
        cost_pc=minutes*rate
        proc_cost_pc+=cost_pc
        proc_detail.append((proc, round(rate,2), round(minutes,2), round(cost_pc,2)))
    total_pc=mat_eur_pc+proc_cost_pc
    rows.append({
        "item": p.get("item_code","?"),
        "qty": int(p.get("qty",1)),
        "grade": grade,
        "family": fam,
        "mass": round(m_kg,4),
        "eur_kg": round(eur_per_kg,4),
        "mat_pc": round(mat_eur_pc,2),
        "proc_pc": round(proc_cost_pc,2),
        "tot_pc": round(total_pc,2),
        "proc_detail": proc_detail
    })

df=pd.DataFrame(rows)

# ---- DOCX opbouwen
doc=Document()
doc.add_heading(f"Offer {project_code} — {client_name}", level=1)
p=doc.add_paragraph()
p.add_run(f"Date: {date.today()}").bold=True
p.add_run(f"    Contact: {client_contact} | {client_email}")

doc.add_heading("Scope", level=2)
doc.add_paragraph(f"Assembly: {assembly.get('name','')} — quantity {assembly.get('qty',1)}.")

doc.add_heading("Cost breakdown (per piece, EUR)", level=2)
table=doc.add_table(rows=1, cols=7)
hdr=table.rows[0].cells
hdr[0].text="Item"; hdr[1].text="Grade"; hdr[2].text="Mass (kg/pc)"
hdr[3].text="€/kg"; hdr[4].text="Material €/pc"; hdr[5].text="Proc €/pc"; hdr[6].text="Total €/pc"

for r in rows:
    row=table.add_row().cells
    row[0].text=str(r["item"])
    row[1].text=str(r["grade"])
    row[2].text=f'{r["mass"]:.4f}'
    row[3].text=f'{r["eur_kg"]:.4f}'
    row[4].text=f'{r["mat_pc"]:.2f}'
    row[5].text=f'{r["proc_pc"]:.2f}'
    row[6].text=f'{r["tot_pc"]:.2f}'

total_value = float((df["tot_pc"] * df["qty"]).sum()) if not df.empty else 0.0
doc.add_paragraph().add_run(f"Assembly total (qty × Total €/pc): € {total_value:,.2f}").bold=True

doc.add_heading("Process detail", level=2)
for r in rows:
    doc.add_heading(r["item"], level=3)
    if not r["proc_detail"]:
        doc.add_paragraph("No process cost lines.")
        continue
    t=doc.add_table(rows=1, cols=4)
    th=t.rows[0].cells
    th[0].text="Process"; th[1].text="Rate (€/min)"; th[2].text="Minutes"; th[3].text="Cost €/pc"
    for (proc, rate, mins, cost) in r["proc_detail"]:
        rw=t.add_row().cells
        rw[0].text=str(proc); rw[1].text=f"{rate:.2f}"; rw[2].text=f"{mins:.2f}"; rw[3].text=f"{cost:.2f}"

doc.add_heading("Assumptions", level=2)
ul = doc.add_paragraph(style=None)
ul.add_run(f"Lead time: {int(lead_weeks)} weeks after order.\n")
ul.add_run("Incoterms: EXW.\n")
ul.add_run("Weld quality: EN ISO 5817 level C.\n")
ul.add_run("Scrap: 3%.\n")
ul.add_run("Prices excl. VAT.")

# ---- Download knop
buf=io.BytesIO()
doc.save(buf)
buf.seek(0)
st.download_button(
    "⬇️ Download offerte.docx",
    data=buf,
    file_name=f"offerte_{project_code}.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

st.caption("Deze DOCX gebruikt python-docx, dus geen apart .docx-template nodig. Alle gegevens komen uit je BOM + prijzen + uurtarieven.")
